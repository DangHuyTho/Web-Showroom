#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Markdown to Word (.docx) - Version 2
Handles Vietnamese characters and markdown formatting
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Calibri'
    return heading

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'E7E6E6')
    p._element.get_or_add_pPr().append(shading)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

def add_table(doc, table_lines):
    headers = [h.strip() for h in table_lines[0].split('|')[1:-1]]
    rows = []
    for line in table_lines[2:]:
        if line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
    
    if not rows:
        return
    
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Headers
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = cell
            for p in table.rows[r_idx + 1].cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)

# Read markdown
with open('BAO_CAO_DATN_PHAN_1.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Create document
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

lines = content.split('\n')
i = 0
in_code = False
code_lines = []

while i < len(lines):
    line = lines[i]
    
    # Headings
    if line.startswith('# '):
        add_heading(doc, line[2:].strip(), 1)
    elif line.startswith('## '):
        add_heading(doc, line[3:].strip(), 2)
    elif line.startswith('### '):
        add_heading(doc, line[4:].strip(), 3)
    elif line.startswith('#### '):
        add_heading(doc, line[5:].strip(), 4)
    elif line.startswith('##### '):
        add_heading(doc, line[6:].strip(), 5)
    
    # Code blocks
    elif line.strip().startswith('```'):
        if not in_code:
            in_code = True
            code_lines = []
        else:
            in_code = False
            if code_lines:
                add_code_block(doc, '\n'.join(code_lines))
            code_lines = []
    elif in_code:
        code_lines.append(line)
    
    # Tables
    elif '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
        table_lines = [line]
        idx = i + 1
        while idx < len(lines) and lines[idx].strip().startswith('|'):
            table_lines.append(lines[idx])
            idx += 1
        add_table(doc, table_lines)
        i = idx - 1
    
    # Bullet lists
    elif line.strip().startswith('- '):
        doc.add_paragraph(line.strip()[2:].strip(), style='List Bullet')
    
    # Regular paragraphs
    elif line.strip():
        p = doc.add_paragraph(line.strip())
        for run in p.runs:
            run.font.name = 'Calibri'
    
    i += 1

# Save document
doc.save('BAO_CAO_DATN_V2_COMPLETE.docx')
print('✓ File mới: BAO_CAO_DATN_V2_COMPLETE.docx đã tạo thành công!')
print('✓ Vị trí: c:\\Users\\ThoDh\\DATN\\BAO_CAO_DATN_V2_COMPLETE.docx')
