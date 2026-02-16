#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Markdown to Word DOCX with proper table support
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys

def markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX with table support"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # Skip empty lines but preserve spacing
        if not line.strip():
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=1)
            for run in heading.runs:
                run.font.name = 'Calibri'
            i += 1
            
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), level=2)
            for run in heading.runs:
                run.font.name = 'Calibri'
            i += 1
            
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), level=3)
            for run in heading.runs:
                run.font.name = 'Calibri'
            i += 1
            
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:].strip(), level=4)
            for run in heading.runs:
                run.font.name = 'Calibri'
            i += 1
        
        # Handle markdown tables
        elif line.strip().startswith('|') and i + 1 < len(lines):
            # Check if next line is separator
            next_line = lines[i + 1].strip()
            if next_line.startswith('|') and all(c in '-| :' for c in next_line):
                # Parse table header
                headers = [h.strip() for h in line.split('|')[1:-1]]
                
                # Collect table rows
                rows = []
                j = i + 2
                while j < len(lines):
                    row_line = lines[j].strip()
                    if not row_line.startswith('|'):
                        break
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    if len(cells) == len(headers):
                        rows.append(cells)
                    j += 1
                
                # Create table
                if rows:
                    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Add header row
                    header_cells = table.rows[0].cells
                    for col_idx, header in enumerate(headers):
                        cell = header_cells[col_idx]
                        cell.text = header
                        
                        # Style header
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(11)
                                run.font.name = 'Calibri'
                    
                    # Add data rows
                    for row_idx, row in enumerate(rows):
                        row_cells = table.rows[row_idx + 1].cells
                        for col_idx, cell_data in enumerate(row):
                            if col_idx < len(row_cells):
                                row_cells[col_idx].text = cell_data
                                for paragraph in row_cells[col_idx].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(10)
                                        run.font.name = 'Calibri'
                
                i = j
            else:
                # Not a table, add as regular paragraph
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.name = 'Calibri'
                i += 1
        
        # Handle horizontal rules
        elif line.strip() == '---':
            # Add empty paragraph for spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
        
        # Handle code blocks
        elif line.strip().startswith('```'):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            
            if code_lines:
                # Add code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.25)
                
                # Background color
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'E7E6E6')
                p._element.get_or_add_pPr().append(shading_elm)
                
                # Code text
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            
            # Skip closing ```
            if i < len(lines) and lines[i].strip().startswith('```'):
                i += 1
        
        # Handle lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            # Determine indent level
            indent_level = (len(line) - len(line.lstrip(' '))) // 4
            list_text = line.strip()[2:].strip()
            
            # Add list item
            p = doc.add_paragraph(list_text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            i += 1
        
        # Regular text with formatting
        else:
            # Handle bold and italic
            text = line.strip()
            p = doc.add_paragraph()
            
            # Simple regex-based formatting
            import re
            
            # Split by markdown formatting
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_)', text)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Bold
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                    run.font.name = 'Calibri'
                elif part.startswith('*') and part.endswith('*'):
                    # Italic
                    run = p.add_run(part[1:-1])
                    run.font.italic = True
                    run.font.name = 'Calibri'
                elif part.startswith('__') and part.endswith('__'):
                    # Bold
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                    run.font.name = 'Calibri'
                elif part.startswith('_') and part.endswith('_'):
                    # Italic
                    run = p.add_run(part[1:-1])
                    run.font.italic = True
                    run.font.name = 'Calibri'
                elif part:
                    # Normal text
                    run = p.add_run(part)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            
            i += 1
    
    # Save document
    doc.save(docx_file)

if __name__ == '__main__':
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    else:
        md_file = r'c:\Users\ThoDh\DATN\BAO_CAO_DATN_PHAN_1.md'
    
    # Generate output filename
    base_name = os.path.splitext(os.path.basename(md_file))[0]
    docx_file = os.path.join(os.path.dirname(md_file), f'{base_name}.docx')
    
    print(f"Converting: {md_file}")
    print(f"Output: {docx_file}")
    
    if os.path.exists(md_file):
        markdown_to_docx(md_file, docx_file)
        print(f"\n✅ Conversion successful!")
        print(f"📄 File: {os.path.basename(docx_file)}")
        print(f"📍 Location: {os.path.abspath(docx_file)}")
        
        size_mb = os.path.getsize(docx_file) / (1024 * 1024)
        print(f"📊 File size: {size_mb:.2f} MB")
    else:
        print(f"❌ Error: File not found - {md_file}")
