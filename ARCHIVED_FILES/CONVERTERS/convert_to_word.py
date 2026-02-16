#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Markdown report to Word (.docx) format
Handles Vietnamese characters and formatting properly
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_heading_with_style(doc, text, level):
    """Add heading with proper styling"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add Vietnamese font for proper display
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(12 + (4 - level) * 2)
    
    return heading

def add_table_from_markdown(doc, lines, start_idx):
    """Parse and add table from markdown format"""
    header_line = lines[start_idx].strip()
    
    # Parse header
    headers = [h.strip() for h in header_line.split('|')[1:-1]]
    
    # Find table rows (skip separator line at start_idx + 1)
    rows_data = []
    idx = start_idx + 2
    while idx < len(lines):
        line = lines[idx].strip()
        # Stop if line doesn't start with | or is empty
        if not line or not line.startswith('|'):
            break
        cells = [c.strip() for c in line.split('|')[1:-1]]
        # Only add if number of cells matches headers
        if len(cells) == len(headers):
            rows_data.append(cells)
        idx += 1
    
    # Create table with proper dimensions
    if rows_data:
        table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add headers with background color
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            
            # Set cell background to light gray
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D3D3D3')
            cell._element.get_or_add_tcPr().append(shading_elm)
            
            # Format header text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
        
        # Add rows
        for row_idx, row_data in enumerate(rows_data):
            row_cells = table.rows[row_idx + 1].cells
            for cell_idx, cell_data in enumerate(row_data):
                if cell_idx < len(row_cells):
                    row_cells[cell_idx].text = cell_data
                    for paragraph in row_cells[cell_idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(10)
    
    return idx

def is_table_separator(line):
    """Check if line is a valid markdown table separator"""
    line = line.strip()
    if not line.startswith('|'):
        return False
    parts = [p.strip() for p in line.split('|')[1:-1]]
    # Check if all parts contain dashes
    return all(all(c in '-: ' for c in p) for p in parts)
    """Add code block with monospace font"""
    paragraph = doc.add_paragraph()
    
    # Set paragraph style
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Inches(0.25)
    paragraph_format.right_indent = Inches(0.25)
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    
    # Add background color
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E7E6E6')
    paragraph._element.get_or_add_pPr().append(shading_elm)
    
    # Add code with monospace font
    run = paragraph.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    return paragraph

def markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    i = 0
    in_code_block = False
    code_block_content = []
    current_code_lang = ''
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                current_code_lang = line.strip()[3:].strip()
                code_block_content = []
            else:
                in_code_block = False
                if code_block_content:
                    add_code_block(doc, '\n'.join(code_block_content), current_code_lang)
                code_block_content = []
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            add_heading_with_style(doc, line[2:].strip(), level=1)
        elif line.startswith('## '):
            add_heading_with_style(doc, line[3:].strip(), level=2)
        elif line.startswith('### '):
            add_heading_with_style(doc, line[4:].strip(), level=3)
        elif line.startswith('#### '):
            add_heading_with_style(doc, line[5:].strip(), level=4)
        
        # Handle tables
        elif '|' in line and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            i = add_table_from_markdown(doc, lines, i) - 1
        
        # Handle unordered lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            paragraph = doc.add_paragraph(line.strip()[2:].strip(), style='List Bullet')
            for run in paragraph.runs:
                run.font.name = 'Calibri'
        
        # Handle dividers
        elif line.strip() == '---':
            doc.add_paragraph()
        
        # Handle regular paragraphs
        elif line.strip() and not line.startswith('```'):
            paragraph = doc.add_paragraph(line.strip())
            # Set font for Vietnamese support
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✓ Document converted successfully!")
    print(f"✓ Saved to: {docx_file}")

if __name__ == '__main__':
    import sys
    import os
    
    # Use argument if provided, otherwise default
    if len(sys.argv) > 1:
        md_input = sys.argv[1]
    else:
        md_input = r'c:\Users\ThoDh\DATN\BAO_CAO_DATN_PHAN_1.md'
    
    # Generate output filename from input
    base_name = os.path.splitext(os.path.basename(md_input))[0]
    docx_output = os.path.join(os.path.dirname(md_input), f'{base_name}.docx')
    
    if os.path.exists(md_input):
        markdown_to_docx(md_input, docx_output)
        print(f"\n📄 File: {os.path.basename(docx_output)}")
        print(f"📍 Location: {docx_output}")
        
        # Get file size
        size_mb = os.path.getsize(docx_output) / (1024 * 1024)
        print(f"📊 File size: {size_mb:.2f} MB")
    else:
        print(f"❌ Error: File not found - {md_input}")
